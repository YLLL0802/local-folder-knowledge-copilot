from __future__ import annotations

import argparse
import hashlib
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from .permissions import (
    allowed_roles_for_department,
    infer_department,
    sensitivity_for_department,
)
from .schemas import DocumentMetadata, LoadedDocument


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}
EXCLUDED_DIRS = {
    ".git",
    ".venv",
    "venv",
    "env",
    "node_modules",
    "__pycache__",
    ".qdrant",
    ".cache",
    "cache",
    ".pytest_cache",
}
TEMP_FILE_PREFIXES = ("~$", ".~")


def load_documents(root_path: str | Path, max_file_size_mb: int = 10) -> list[LoadedDocument]:
    """Load supported documents from a folder into text plus metadata objects."""

    root = Path(root_path).expanduser().resolve()
    if not root.exists():
        raise FileNotFoundError(f"Document root does not exist: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"Document root is not a directory: {root}")

    max_bytes = max_file_size_mb * 1024 * 1024
    documents: list[LoadedDocument] = []
    for file_path in iter_supported_files(root, max_bytes=max_bytes):
        text = read_text(file_path)
        if not text.strip():
            continue
        documents.append(build_loaded_document(file_path, root, text))
    return documents


def iter_supported_files(root: Path, max_bytes: int) -> Iterable[Path]:
    """Yield supported files under root while skipping excluded paths and large files."""

    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if _is_excluded(path, root):
            continue
        if path.name.startswith(TEMP_FILE_PREFIXES):
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if path.stat().st_size > max_bytes:
            continue
        yield path


def read_text(path: Path) -> str:
    """Extract plain text from a supported Markdown, text, PDF, or DOCX file."""

    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file extension: {path.suffix}")


def build_loaded_document(path: Path, root: Path, text: str) -> LoadedDocument:
    """Build a LoadedDocument and infer enterprise-style metadata from its path."""

    stat = path.stat()
    checksum = hashlib.sha256(text.encode("utf-8")).hexdigest()
    modified_at = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc)
    relative_path = path.relative_to(root).as_posix()
    department = infer_department(path, root)
    doc_id_seed = f"{path.resolve()}:{stat.st_mtime_ns}"
    doc_id = hashlib.sha256(doc_id_seed.encode("utf-8")).hexdigest()[:24]

    metadata = DocumentMetadata(
        doc_id=doc_id,
        source_path=str(path.resolve()),
        relative_path=relative_path,
        file_name=path.name,
        file_type=path.suffix.lower().lstrip("."),
        modified_at=modified_at,
        department=department,
        sensitivity=sensitivity_for_department(department),
        allowed_roles=allowed_roles_for_department(department),
        checksum=checksum,
        size_bytes=stat.st_size,
    )
    return LoadedDocument(text=text, metadata=metadata)


def _is_excluded(path: Path, root: Path) -> bool:
    """Return true when a path sits under an ignored directory."""

    try:
        relative_parts = path.relative_to(root).parts
    except ValueError:
        return True
    return any(part in EXCLUDED_DIRS for part in relative_parts)


def _read_pdf(path: Path) -> str:
    """Extract text from all pages of a PDF file."""

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    """Extract paragraph text from a DOCX file."""

    from docx import Document

    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _main() -> None:
    """Run a small CLI preview for loaded documents and inferred metadata."""

    parser = argparse.ArgumentParser(description="Preview loaded local OneDrive documents.")
    parser.add_argument("--root", default="data/sample_docs", help="Folder to scan.")
    parser.add_argument("--limit", type=int, default=20, help="Maximum rows to print.")
    args = parser.parse_args()

    docs = load_documents(args.root)
    print(f"Loaded {len(docs)} documents from {Path(args.root).resolve()}")
    for doc in docs[: args.limit]:
        metadata = doc.metadata
        print(
            f"- {metadata.relative_path} | department={metadata.department} | "
            f"sensitivity={metadata.sensitivity} | roles={','.join(metadata.allowed_roles)}"
        )
        print(f"  {doc.preview}")


if __name__ == "__main__":
    _main()
